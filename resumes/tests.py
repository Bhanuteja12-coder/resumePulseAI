import io

from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from rest_framework import status
from rest_framework.test import APITestCase

from accounts.models import User


class ResumeApiTests(APITestCase):
    def setUp(self):
        self.user = User.objects.create_user(email='resumeuser@example.com', password='ResumePass123!')
        self.upload_url = reverse('resume-upload')
        self.login_url = reverse('token_obtain_pair')

        response = self.client.post(self.login_url, {'email': self.user.email, 'password': 'ResumePass123!'}, format='json')
        self.access_token = response.data['access']

    def authenticate(self):
        self.client.credentials(HTTP_AUTHORIZATION=f'Bearer {self.access_token}')

    def test_resume_upload_requires_auth(self):
        file_data = SimpleUploadedFile('resume.pdf', b'PDF file content', content_type='application/pdf')
        response = self.client.post(self.upload_url, {'file': file_data}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_401_UNAUTHORIZED)

    def test_resume_upload_accepts_pdf_and_extracts_text(self):
        import fitz

        self.authenticate()
        pdf_buffer = fitz.open()
        page = pdf_buffer.new_page()
        page.insert_text((72, 72), 'Hello PDF Resume')
        file_data = SimpleUploadedFile('resume.pdf', pdf_buffer.write(), content_type='application/pdf')

        response = self.client.post(self.upload_url, {'file': file_data}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertEqual(response.data['user'], self.user.id)
        self.assertIn('Hello PDF Resume', response.data['extracted_text'])

    def test_resume_upload_accepts_docx_and_extracts_text(self):
        from docx import Document

        self.authenticate()
        document = Document()
        document.add_paragraph('Hello DOCX Resume')
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        file_data = SimpleUploadedFile(
            'resume.docx',
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        response = self.client.post(self.upload_url, {'file': file_data}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('Hello DOCX Resume', response.data['extracted_text'])

    def test_resume_upload_rejects_invalid_extension(self):
        self.authenticate()
        bad_file = SimpleUploadedFile('resume.txt', b'text content', content_type='text/plain')
        response = self.client.post(self.upload_url, {'file': bad_file}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn('file', response.data)

    def test_resume_upload_extracts_skill_keywords(self):
        self.authenticate()
        document = self._create_docx_file(
            'Python developer with experience in machine learning and data analysis.'
        )
        response = self.client.post(self.upload_url, {'file': document}, format='multipart')
        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('python', response.data['skill_keywords'])
        self.assertIn('machine', response.data['skill_keywords'])
        self.assertNotIn('with', response.data['skill_keywords'])

    def test_get_resume_list_returns_uploaded_resume(self):
        self.authenticate()
        document = self._create_docx_file('Hello DOCX Resume')
        upload_response = self.client.post(self.upload_url, {'file': document}, format='multipart')
        self.assertEqual(upload_response.status_code, status.HTTP_201_CREATED)

        response = self.client.get('/api/resumes/')
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data), 1)
        self.assertIn('Hello DOCX Resume', response.data[0]['extracted_text'])

    def test_get_resume_detail_returns_extracted_text(self):
        self.authenticate()
        document = self._create_docx_file('Hello DOCX Resume Detail')
        upload_response = self.client.post(self.upload_url, {'file': document}, format='multipart')
        self.assertEqual(upload_response.status_code, status.HTTP_201_CREATED)
        resume_id = upload_response.data['id']

        response = self.client.get(f'/api/resumes/{resume_id}/')
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertIn('Hello DOCX Resume Detail', response.data['extracted_text'])

    def test_resume_search_by_skill_keyword(self):
        self.authenticate()
        document = self._create_docx_file(
            'Experienced Python engineer with Django and machine learning expertise.'
        )
        upload_response = self.client.post(self.upload_url, {'file': document}, format='multipart')
        self.assertEqual(upload_response.status_code, status.HTTP_201_CREATED)

        response = self.client.get('/api/resumes/search/', {'skill': 'python'})
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data), 1)
        self.assertEqual(response.data[0]['id'], upload_response.data['id'])

        response = self.client.get('/api/resumes/search/', {'skill': 'sql'})
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response.data), 0)

    def test_resume_analysis_creates_match_report(self):
        from jobs.models import JobDescription

        self.authenticate()

        document = self._create_docx_file(
            'Experienced Python engineer with Django and machine learning expertise.'
        )
        upload_response = self.client.post(self.upload_url, {'file': document}, format='multipart')
        self.assertEqual(upload_response.status_code, status.HTTP_201_CREATED)
        resume_id = upload_response.data['id']

        jd = JobDescription.objects.create(
            title='Machine Learning Engineer',
            company='AI Labs',
            location='Remote',
            description='Looking for a Python developer experienced with Django and machine learning.',
        )

        analysis_url = reverse('resume-analyze')
        response = self.client.post(
            analysis_url,
            {'resume_id': resume_id, 'job_description_id': jd.id},
            format='json'
        )

        self.assertEqual(response.status_code, status.HTTP_201_CREATED)
        self.assertIn('match_score', response.data)
        self.assertIn('match_percent', response.data)
        self.assertIn('gap_analysis', response.data)
        self.assertIsInstance(response.data['gap_analysis'], dict)
        self.assertIn('skills', response.data['gap_analysis'])
        self.assertIn('tools', response.data['gap_analysis'])
        self.assertIn('experience', response.data['gap_analysis'])
        self.assertNotIn('build', response.data['gap_analysis']['skills'])
        self.assertNotIn('great', response.data['gap_analysis']['skills'])
        self.assertNotIn('products', response.data['gap_analysis']['skills'])
        self.assertGreater(response.data['match_score'], 0.0)
        self.assertGreater(response.data['match_percent'], 0.0)

    def test_report_list_is_paginated_and_filtered_to_user(self):
        from jobs.models import JobDescription
        from resumes.models import AnalysisReport

        self.authenticate()

        document = self._create_docx_file('Resume for report test')
        upload_response = self.client.post(self.upload_url, {'file': document}, format='multipart')
        self.assertEqual(upload_response.status_code, status.HTTP_201_CREATED)
        resume_id = upload_response.data['id']

        jd = JobDescription.objects.create(
            title='Backend Developer',
            company='Tech Co',
            location='Remote',
            description='Python Django REST API developer required.',
        )

        for i in range(12):
            AnalysisReport.objects.create(
                resume_id=resume_id,
                job_description=jd,
                match_score=0.5,
                match_percent=50.0,
                ai_suggestions=['Suggestion'],
            )

        response = self.client.get('/api/reports/')
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data['count'], 12)
        self.assertEqual(len(response.data['results']), 10)

        response_page_2 = self.client.get('/api/reports/', {'page': 2})
        self.assertEqual(response_page_2.status_code, status.HTTP_200_OK)
        self.assertEqual(len(response_page_2.data['results']), 2)

    def test_report_detail_returns_report_data(self):
        from jobs.models import JobDescription
        from resumes.models import AnalysisReport

        self.authenticate()

        document = self._create_docx_file('Resume report detail test')
        upload_response = self.client.post(self.upload_url, {'file': document}, format='multipart')
        self.assertEqual(upload_response.status_code, status.HTTP_201_CREATED)
        resume_id = upload_response.data['id']

        jd = JobDescription.objects.create(
            title='Backend Engineer',
            company='Example Co',
            location='Remote',
            description='Looking for Django and Python experience.',
        )

        report = AnalysisReport.objects.create(
            resume_id=resume_id,
            job_description=jd,
            match_score=0.7,
            match_percent=70.0,
            ai_suggestions=['Suggestion one', 'Suggestion two'],
        )

        response = self.client.get(f'/api/reports/{report.id}/')
        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response.data['id'], report.id)
        self.assertEqual(response.data['match_percent'], 70.0)
        self.assertEqual(response.data['ai_suggestions'], ['Suggestion one', 'Suggestion two'])

    def test_report_end_to_end_analysis_list_and_detail(self):
        from jobs.models import JobDescription

        self.authenticate()

        document = self._create_docx_file('Full end to end report resume test')
        upload_response = self.client.post(self.upload_url, {'file': document}, format='multipart')
        self.assertEqual(upload_response.status_code, status.HTTP_201_CREATED)
        resume_id = upload_response.data['id']

        jd = JobDescription.objects.create(
            title='Django Developer',
            company='ExampleCorp',
            location='Remote',
            description='Looking for Django, Python, and MongoDB experience.',
        )

        analysis_url = reverse('resume-analyze')
        create_response = self.client.post(
            analysis_url,
            {'resume_id': resume_id, 'job_description_id': jd.id},
            format='json'
        )
        self.assertEqual(create_response.status_code, status.HTTP_201_CREATED)
        report_id = create_response.data['id']

        list_response = self.client.get(reverse('report-list'))
        self.assertEqual(list_response.status_code, status.HTTP_200_OK)
        self.assertGreaterEqual(list_response.data['count'], 1)
        self.assertTrue(any(item['id'] == report_id for item in list_response.data['results']))

        detail_response = self.client.get(reverse('report-detail', args=[report_id]))
        self.assertEqual(detail_response.status_code, status.HTTP_200_OK)
        self.assertEqual(detail_response.data['id'], report_id)
        self.assertEqual(detail_response.data['job_description']['id'], jd.id)
        self.assertIsInstance(detail_response.data['ai_suggestions'], list)

    def _create_docx_file(self, text):
        from docx import Document

        document = Document()
        document.add_paragraph(text)
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return SimpleUploadedFile(
            'resume.docx',
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
